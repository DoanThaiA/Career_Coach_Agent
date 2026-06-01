import os
import fitz
import docx
from typing import List
from src.core.logger import get_logger
from langchain_core.documents import Document

logger = get_logger(__name__)

class DocumentExtractService:
    def __init__(self):
        pass

    def process_file(self, file_path: str) -> Document:
       if not os.path.exists(file_path):
            raise FileNotFoundError(f"File không tồn tại: {file_path}")
       
       ext = os.path.splitext(file_path)[1].lower()
       logger.info(f"Processing file {file_path} with extension {ext}")
       
       documents = []
       if ext == '.txt':
            documents = self._extract_txt(file_path)
       elif ext == '.docx':
            documents = self._extract_docx(file_path)
       elif ext == '.pdf':
            documents = self._extract_pdf(file_path)
       else:
            raise ValueError(f"Định dạng {ext} chưa được hỗ trợ (OCR đã bị vô hiệu hóa)")
            
       full_text = "\n".join(doc.page_content for doc in documents if doc.page_content)
            
       return Document(
           page_content=full_text,
           metadata={
               "source": file_path,
               "extension": ext, 
               "num_documents": len(documents),
               "total_length": len(full_text)
           }
       )
       
    def _extract_txt(self, file_path: str) -> List[Document]:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
            return [Document(page_content=text, metadata={"source": file_path})]

    def _extract_docx(self, file_path: str) -> List[Document]:
        doc = docx.Document(file_path)
        doc_list = []
        
        from docx.document import Document as _Document
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        parent_elm = doc.element.body
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                para = Paragraph(child, doc)
                if para.text.strip():
                    doc_list.append(Document(page_content=para.text, metadata={
                        "source": file_path,
                        "type": "paragraph",
                    }))
            elif isinstance(child, CT_Tbl):
                table = Table(child, doc)
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    doc_list.append(Document(page_content="\t".join(row_data), metadata={
                        "source": file_path,
                        "type": "table",
                    }))
        return doc_list

    def _extract_pdf(self, file_path: str) -> List[Document]:
        pdf_doc = fitz.open(file_path)
        doc_list = []
        try:
            for page in pdf_doc:
                text = page.get_text()
                if text.strip():
                    doc_list.append(Document(page_content=text, metadata={
                        "source": file_path,
                        "page": page.number + 1,
                        "type": "page",
                    }))
        finally:
            pdf_doc.close()
        return doc_list